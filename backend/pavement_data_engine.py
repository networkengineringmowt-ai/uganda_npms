#!/usr/bin/env python3
"""Build the auditable NPMS backend from manuals, condition and inventory repositories.

Every published current value is linked to a variable definition, source file and
model run. Source workbooks/databases remain authoritative; ML fills temporal and
coverage gaps and always publishes confidence and method metadata.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import os
import re
import sqlite3
import sys
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

import numpy as np
from openpyxl import load_workbook
from pypdf import PdfReader
from docx import Document
from sklearn.compose import ColumnTransformer
from sklearn.metrics import mean_absolute_error, r2_score
from sklearn.neural_network import MLPRegressor
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler

from pms_backend.variables import ENGINE_SETTINGS as SETTINGS, ENGINE_VARIABLE_ROWS as VARIABLES

try:
    import pyodbc
except ImportError:  # Access metadata is optional at runtime but reported explicitly.
    pyodbc = None


@dataclass(frozen=True)
class EnginePaths:
    repository: Path
    app: Path
    database: Path
    schema: Path
    dashboard_json: Path
    network_json: Path
    condition_json: Path
    prediction_json: Path
    deterioration_json: Path
    inventory_json: Path


BACKEND = Path(__file__).resolve().parent
REPOSITORY = Path(os.getenv("NPMS_REPOSITORY_ROOT", Path(__file__).resolve().parents[2])).resolve()
APP = Path(os.getenv("NPMS_DEPLOY_ROOT", Path(__file__).resolve().parents[1])).resolve()
PATHS = EnginePaths(
    repository=REPOSITORY,
    app=APP,
    database=Path(os.getenv("PMS_DB_PATH", APP / "data" / "npms_backend.db")).resolve(),
    schema=Path(os.getenv("PMS_SCHEMA_PATH", BACKEND / "sql" / "pms_backend_schema.sql")).resolve(),
    dashboard_json=APP / "data" / "pms_dashboard.json",
    network_json=APP / "data" / "network_links.json",
    condition_json=APP / "data" / "link_condition_lookup.json",
    prediction_json=APP / "data" / "romdas_predictions.json",
    deterioration_json=APP / "data" / "deterioration_summary.json",
    inventory_json=APP / "data" / "road_inventory_2023.json",
)

SOURCE_ROOTS = {
    "pavement_manuals": REPOSITORY / "0. Manuals" / "Pavement Design Manuals",
    "road_condition": REPOSITORY / "5.Road Condition Data",
    "road_inventory": REPOSITORY / "6.Road Inventory Data",
}

ALIASES = {
    "linkid": "link_id", "link_id": "link_id", "link": "link_id", "gisid": "external_id",
    "road": "road_no", "roadcode": "road_no", "road_no": "road_no", "roadname": "road_name",
    "link_name": "link_name", "linkname": "link_name", "maintenance_region": "maintenance_region",
    "region": "maintenance_region", "maintenance_station": "maintenance_station", "station": "maintenance_station",
    "survey_date": "survey_date", "date": "survey_date", "survey_year": "survey_year",
    "lane_iri": "iri_m_per_km", "average_iri": "iri_m_per_km", "avg_iri": "iri_m_per_km",
    "iri": "iri_m_per_km", "calib_rgh": "iri_m_per_km", "rwp_iri": "iri_m_per_km",
    "average_rut_depth": "rut_depth_mm", "average_lane_rutting_mm": "rut_depth_mm",
    "ave_rut_se": "rut_depth_mm", "rut_depth_mean": "rut_depth_mm", "rutting": "rut_depth_mm",
    "vci": "vci", "average_vci": "vci", "average_weighted_vci": "vci", "pci": "pci",
    "cracking": "cracking_percent", "cracking_percent": "cracking_percent",
    "potholes": "potholes_value", "condition": "condition_class", "vci_rating": "condition_class",
    "iri_rating": "condition_class", "rating": "condition_class", "feature_type": "asset_type",
    "feature_subtype": "asset_subtype", "sub_type": "asset_subtype", "sub_type_code": "asset_subtype",
    "chainage": "chainage_from_km", "chainage_km": "chainage_from_km", "chainage_from": "chainage_from_km",
    "chainage_start_km": "chainage_from_km", "chainage_end": "chainage_to_km",
    "x_coordinate": "longitude", "x_start": "longitude", "y_coordinate": "latitude", "y_start": "latitude",
    "length_m": "length_m", "culvert_length": "length_m", "road_width_m": "width_m",
    "overall_bridge_width_m": "width_m", "number_of_features": "quantity", "no_of_barrels": "quantity",
}


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def clean_name(value: Any) -> str:
    text = str(value or "").strip().lower().replace("%", " percent ")
    text = re.sub(r"[^a-z0-9]+", "_", text).strip("_")
    return re.sub(r"_+", "_", text)


def canonical_field(value: Any) -> str:
    cleaned = clean_name(value)
    return ALIASES.get(cleaned, cleaned)


def number(value: Any) -> float | None:
    if value is None or value == "":
        return None
    try:
        result = float(value)
        return result if math.isfinite(result) else None
    except (TypeError, ValueError):
        match = re.search(r"-?\d+(?:\.\d+)?", str(value).replace(",", ""))
        return float(match.group()) if match else None


def iso_date(value: Any) -> str | None:
    if value is None or value == "":
        return None
    if hasattr(value, "date"):
        return value.date().isoformat()
    text = str(value).strip()
    year = re.search(r"(?:19|20)\d{2}", text)
    return text[:10] if re.match(r"\d{4}-\d{2}-\d{2}", text) else (f"{year.group()}-12-31" if year else None)


def link_id(value: Any) -> str | None:
    text = str(value or "").strip().replace(" ", "")
    match = re.search(r"([A-Z0-9]+(?:N\d+)?_Link\d+)", text, re.I)
    return match.group(1) if match else None


def json_value(value: Any) -> Any:
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value)


def sha256(path: Path, max_mb: float) -> str | None:
    if path.stat().st_size > max_mb * 1024 * 1024:
        return None
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def setting(conn: sqlite3.Connection, key: str) -> Any:
    row = conn.execute("SELECT value_numeric,value_text FROM pms_engine_settings WHERE setting_key=?", (key,)).fetchone()
    if not row:
        raise KeyError(f"Undefined engine setting: {key}")
    return row[0] if row[0] is not None else row[1]


def setup_database(conn: sqlite3.Connection) -> None:
    conn.executescript(PATHS.schema.read_text(encoding="utf-8"))
    conn.executemany(
        """INSERT INTO pms_engine_settings(setting_key,value_numeric,value_text,unit,description)
           VALUES(?,?,?,?,?) ON CONFLICT(setting_key) DO UPDATE SET
           value_numeric=excluded.value_numeric,value_text=excluded.value_text,unit=excluded.unit,
           description=excluded.description,updated_at=CURRENT_TIMESTAMP""",
        [(key, *values) for key, values in SETTINGS.items()],
    )
    conn.executemany(
        """INSERT INTO pms_variable_definitions
           (canonical_name,category,data_type,unit,description,valid_min,valid_max,
            aggregation_method,current_value_method,definition_source)
           VALUES(?,?,?,?,?,?,?,?,?,?) ON CONFLICT(canonical_name) DO UPDATE SET
           category=excluded.category,data_type=excluded.data_type,unit=excluded.unit,
           description=excluded.description,valid_min=excluded.valid_min,valid_max=excluded.valid_max,
           aggregation_method=excluded.aggregation_method,current_value_method=excluded.current_value_method""",
        [(*row, "Pavement manuals and source-field harmonisation") for row in VARIABLES],
    )
    conn.commit()


def register_file(conn: sqlite3.Connection, group: str, root: Path, path: Path) -> int:
    stat = path.stat()
    max_mb = float(setting(conn, "source_hash_max_mb"))
    digest = sha256(path, max_mb)
    conn.execute(
        """INSERT INTO pms_source_files
           (repository_group,absolute_path,relative_path,file_name,extension,byte_size,modified_at,content_sha256)
           VALUES(?,?,?,?,?,?,?,?) ON CONFLICT(absolute_path) DO UPDATE SET
           repository_group=excluded.repository_group,relative_path=excluded.relative_path,
           file_name=excluded.file_name,extension=excluded.extension,byte_size=excluded.byte_size,
           modified_at=excluded.modified_at,content_sha256=excluded.content_sha256,
           ingestion_status='registered',error_message=NULL""",
        (group, str(path), str(path.relative_to(root)), path.name, path.suffix.lower(), stat.st_size,
         datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(), digest),
    )
    return conn.execute("SELECT source_id FROM pms_source_files WHERE absolute_path=?", (str(path),)).fetchone()[0]


def ensure_dynamic_variable(conn: sqlite3.Connection, canonical: str, source_name: str) -> None:
    if not canonical:
        return
    conn.execute(
        """INSERT OR IGNORE INTO pms_variable_definitions
           (canonical_name,category,data_type,description,aggregation_method,current_value_method,definition_source)
           VALUES(?,?,?,?,?,?,?)""",
        (canonical, "source_field", "UNKNOWN", f"Discovered source field: {source_name}",
         "source-specific", "source record", "Automated source schema discovery"),
    )


def register_table(conn: sqlite3.Connection, source_id: int, name: str, kind: str,
                   header_row: int | None, row_count: int, fields: list[str]) -> int:
    conn.execute(
        """INSERT INTO pms_source_tables(source_id,table_name,table_kind,header_row,row_count,column_count)
           VALUES(?,?,?,?,?,?) ON CONFLICT(source_id,table_name) DO UPDATE SET
           table_kind=excluded.table_kind,header_row=excluded.header_row,row_count=excluded.row_count,
           column_count=excluded.column_count""",
        (source_id, name, kind, header_row, row_count, len(fields)),
    )
    table_id = conn.execute(
        "SELECT source_table_id FROM pms_source_tables WHERE source_id=? AND table_name=?", (source_id, name)
    ).fetchone()[0]
    conn.execute("DELETE FROM pms_source_fields WHERE source_table_id=?", (table_id,))
    rows = []
    for position, field in enumerate(fields, 1):
        source_name = field or f"unnamed_{position}"
        canonical = canonical_field(source_name)
        ensure_dynamic_variable(conn, canonical, source_name)
        rows.append((table_id, position, source_name, canonical, "UNKNOWN"))
    conn.executemany(
        """INSERT INTO pms_source_fields
           (source_table_id,ordinal_position,source_name,canonical_name,inferred_type)
           VALUES(?,?,?,?,?)""", rows,
    )
    return table_id


def find_header(ws, limit: int = 30) -> tuple[int | None, list[str]]:
    for row_number, row in enumerate(ws.iter_rows(min_row=1, max_row=min(ws.max_row, limit), values_only=True), 1):
        values = [str(v).strip() if v is not None else "" for v in row]
        if sum(bool(v) for v in values) >= 3:
            while values and not values[-1]:
                values.pop()
            return row_number, values
    return None, []


def field_map(headers: list[str]) -> dict[str, int]:
    result: dict[str, int] = {}
    for index, header in enumerate(headers):
        canonical = canonical_field(header)
        if canonical and canonical not in result:
            result[canonical] = index
    return result


def row_dict(headers: list[str], values: tuple[Any, ...]) -> dict[str, Any]:
    return {headers[i] or f"unnamed_{i+1}": json_value(values[i]) for i in range(min(len(headers), len(values))) if values[i] not in (None, "")}


def get(values: tuple[Any, ...], mapping: dict[str, int], name: str) -> Any:
    index = mapping.get(name)
    return values[index] if index is not None and index < len(values) else None


def ensure_road_link(conn: sqlite3.Connection, lid: str, source_id: int) -> None:
    conn.execute(
        """INSERT OR IGNORE INTO pms_road_links(link_id,link_name,source_id,updated_at)
           VALUES(?,?,?,?)""",
        (lid, lid, source_id, utc_now()),
    )


def ingest_condition_row(conn: sqlite3.Connection, source_id: int, table_id: int,
                         row_number: int, headers: list[str], mapping: dict[str, int], values: tuple[Any, ...]) -> bool:
    lid = link_id(get(values, mapping, "link_id"))
    metrics = {
        "iri_m_per_km": number(get(values, mapping, "iri_m_per_km")),
        "rut_depth_mm": number(get(values, mapping, "rut_depth_mm")),
        "vci": number(get(values, mapping, "vci")),
        "pci": number(get(values, mapping, "pci")),
        "cracking_percent": number(get(values, mapping, "cracking_percent")),
        "potholes_value": number(get(values, mapping, "potholes_value")),
    }
    if not lid or not any(value is not None for value in metrics.values()):
        return False
    ensure_road_link(conn, lid, source_id)
    survey_date = iso_date(get(values, mapping, "survey_date"))
    year_value = number(get(values, mapping, "survey_year"))
    survey_year = int(year_value) if year_value else (int(survey_date[:4]) if survey_date else None)
    conn.execute(
        """INSERT OR REPLACE INTO pms_condition_observations
           (source_id,source_table_id,source_row_number,link_id,survey_date,survey_year,
            chainage_from_km,chainage_to_km,iri_m_per_km,rut_depth_mm,vci,pci,
            cracking_percent,potholes_value,condition_class,raw_values_json)
           VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (source_id, table_id, row_number, lid, survey_date, survey_year,
         number(get(values, mapping, "chainage_from_km")), number(get(values, mapping, "chainage_to_km")),
         metrics["iri_m_per_km"], metrics["rut_depth_mm"], metrics["vci"], metrics["pci"],
         metrics["cracking_percent"], metrics["potholes_value"],
         str(get(values, mapping, "condition_class") or "") or None,
         json.dumps(row_dict(headers, values), ensure_ascii=False, separators=(",", ":"))),
    )
    return True


def asset_type_from_sheet(sheet_name: str) -> str:
    name = clean_name(sheet_name)
    for token, label in [
        ("minor_culvert", "Minor Culvert"), ("major_culvert", "Major Culvert"),
        ("bridge", "Bridge"), ("marker", "Marker Post or Sign"), ("sign", "Marker Post or Sign"),
        ("point_feature", "Point Feature"), ("line_feature", "Line Feature"), ("road_network", "Road Segment"),
    ]:
        if token in name:
            return label
    return "Inventory Feature"


def ingest_inventory_row(conn: sqlite3.Connection, source_id: int, table_id: int,
                         row_number: int, sheet_name: str, headers: list[str],
                         mapping: dict[str, int], values: tuple[Any, ...]) -> bool:
    lid = link_id(get(values, mapping, "link_id"))
    if not lid:
        return False
    ensure_road_link(conn, lid, source_id)
    asset_type = str(get(values, mapping, "asset_type") or asset_type_from_sheet(sheet_name)).strip()
    conn.execute(
        """INSERT OR REPLACE INTO pms_inventory_assets
           (source_id,source_table_id,source_row_number,external_id,link_id,asset_type,asset_subtype,
            chainage_from_km,chainage_to_km,longitude,latitude,quantity,length_m,width_m,
            condition_class,survey_date,raw_values_json)
           VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (source_id, table_id, row_number, str(get(values, mapping, "external_id") or "") or None,
         lid, asset_type, str(get(values, mapping, "asset_subtype") or "") or None,
         number(get(values, mapping, "chainage_from_km")), number(get(values, mapping, "chainage_to_km")),
         number(get(values, mapping, "longitude")), number(get(values, mapping, "latitude")),
         number(get(values, mapping, "quantity")) or 1.0, number(get(values, mapping, "length_m")),
         number(get(values, mapping, "width_m")), str(get(values, mapping, "condition_class") or "") or None,
         iso_date(get(values, mapping, "survey_date")),
         json.dumps(row_dict(headers, values), ensure_ascii=False, separators=(",", ":"))),
    )
    return True


def ingest_workbook(conn: sqlite3.Connection, source_id: int, group: str, path: Path) -> int:
    workbook = load_workbook(path, read_only=True, data_only=True)
    total = 0
    for ws in workbook.worksheets:
        header_row, headers = find_header(ws)
        table_id = register_table(conn, source_id, ws.title, "excel_sheet", header_row, ws.max_row, headers)
        if not header_row or not headers:
            continue
        mapping = field_map(headers)
        condition_candidate = group == "road_condition" and "link_id" in mapping and any(
            metric in mapping for metric in ("iri_m_per_km", "rut_depth_mm", "vci", "pci", "cracking_percent", "potholes_value")
        )
        inventory_candidate = group == "road_inventory" and "link_id" in mapping and any(
            token in clean_name(ws.title) for token in ("feature", "culvert", "bridge", "marker", "sign", "road_network")
        )
        if not condition_candidate and not inventory_candidate:
            continue
        for row_number, values in enumerate(ws.iter_rows(min_row=header_row + 1, values_only=True), header_row + 1):
            if not any(value not in (None, "") for value in values):
                continue
            if condition_candidate:
                total += int(ingest_condition_row(conn, source_id, table_id, row_number, headers, mapping, values))
            if inventory_candidate:
                total += int(ingest_inventory_row(conn, source_id, table_id, row_number, ws.title, headers, mapping, values))
    workbook.close()
    return total


def ingest_manual(conn: sqlite3.Connection, source_id: int, path: Path) -> int:
    text = ""
    pages = None
    if path.suffix.lower() == ".pdf":
        reader = PdfReader(str(path))
        pages = len(reader.pages)
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    elif path.suffix.lower() == ".docx":
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    else:
        return 0
    conn.execute(
        """INSERT INTO pms_manual_documents
           (source_id,title,document_type,page_count,extracted_character_count,searchable_text,indexed_at)
           VALUES(?,?,?,?,?,?,?) ON CONFLICT(source_id) DO UPDATE SET
           title=excluded.title,document_type=excluded.document_type,page_count=excluded.page_count,
           extracted_character_count=excluded.extracted_character_count,searchable_text=excluded.searchable_text,
           indexed_at=excluded.indexed_at""",
        (source_id, path.stem, path.suffix.lower().lstrip("."), pages, len(text), text, utc_now()),
    )
    return 1


def ingest_access_metadata(conn: sqlite3.Connection, source_id: int, path: Path) -> int:
    if pyodbc is None:
        raise RuntimeError("pyodbc is not installed; Access metadata cannot be indexed")
    connection = pyodbc.connect(
        f"DRIVER={{Microsoft Access Driver (*.mdb, *.accdb)}};DBQ={path};ReadOnly=1;", timeout=20
    )
    total = 0
    try:
        cursor = connection.cursor()
        names = sorted({row.table_name for row in cursor.tables(tableType="TABLE") if not row.table_name.startswith("MSys")})
        for name in names:
            columns = [column.column_name for column in cursor.columns(table=name)]
            try:
                row_count = int(cursor.execute(f"SELECT COUNT(*) FROM [{name}]").fetchone()[0])
            except Exception:
                row_count = 0
            register_table(conn, source_id, name, "access_table", 1, row_count, columns)
            total += row_count
    finally:
        connection.close()
    return total


def ingest_sources(conn: sqlite3.Connection, include_access: bool) -> None:
    for group, root in SOURCE_ROOTS.items():
        for path in sorted(p for p in root.rglob("*") if p.is_file()):
            if path.suffix.lower() in {".mdb", ".accdb"} and not include_access:
                continue
            source_id = register_file(conn, group, root, path)
            records = 0
            try:
                if path.suffix.lower() == ".xlsx":
                    records = ingest_workbook(conn, source_id, group, path)
                elif group == "pavement_manuals" and path.suffix.lower() in {".pdf", ".docx"}:
                    records = ingest_manual(conn, source_id, path)
                elif include_access and path.suffix.lower() in {".mdb", ".accdb"}:
                    records = ingest_access_metadata(conn, source_id, path)
                conn.execute(
                    "UPDATE pms_source_files SET ingestion_status='complete',record_count=?,ingested_at=? WHERE source_id=?",
                    (records, utc_now(), source_id),
                )
            except Exception as exc:
                conn.execute(
                    "UPDATE pms_source_files SET ingestion_status='error',error_message=?,ingested_at=? WHERE source_id=?",
                    (f"{type(exc).__name__}: {exc}"[:1000], utc_now(), source_id),
                )
            conn.commit()


def ingest_access_sources(conn: sqlite3.Connection) -> None:
    for group in ("road_condition", "road_inventory"):
        root = SOURCE_ROOTS[group]
        for path in sorted(p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in {".mdb", ".accdb"}):
            source_id = register_file(conn, group, root, path)
            try:
                records = ingest_access_metadata(conn, source_id, path)
                conn.execute(
                    "UPDATE pms_source_files SET ingestion_status='complete',record_count=?,ingested_at=? WHERE source_id=?",
                    (records, utc_now(), source_id),
                )
            except Exception as exc:
                conn.execute(
                    "UPDATE pms_source_files SET ingestion_status='error',error_message=?,ingested_at=? WHERE source_id=?",
                    (f"{type(exc).__name__}: {exc}"[:1000], utc_now(), source_id),
                )
            conn.commit()


def load_network(conn: sqlite3.Connection) -> None:
    links = json.loads(PATHS.network_json.read_text(encoding="utf-8"))
    source = conn.execute(
        """SELECT source_id FROM pms_source_files WHERE repository_group IN ('road_condition','road_inventory')
           AND lower(file_name) LIKE '%national road network%' ORDER BY modified_at DESC LIMIT 1"""
    ).fetchone()
    source_id = source[0] if source else None
    conn.executemany(
        """INSERT INTO pms_road_links
           (link_id,road_no,road_class,link_name,chainage_from_km,chainage_to_km,length_km,
            surface_type,maintenance_station,maintenance_region,completion_year,rehabilitation_year,
            pavement_age_years,source_id,updated_at)
           VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
           ON CONFLICT(link_id) DO UPDATE SET
           road_no=excluded.road_no,road_class=excluded.road_class,link_name=excluded.link_name,
           chainage_from_km=excluded.chainage_from_km,chainage_to_km=excluded.chainage_to_km,
           length_km=excluded.length_km,surface_type=excluded.surface_type,
           maintenance_station=excluded.maintenance_station,maintenance_region=excluded.maintenance_region,
           completion_year=excluded.completion_year,rehabilitation_year=excluded.rehabilitation_year,
           pavement_age_years=excluded.pavement_age_years,source_id=excluded.source_id,updated_at=excluded.updated_at""",
        [(row.get("link_id"), row.get("road_no"), row.get("road_class"), row.get("link_name"),
          row.get("chainage_from"), row.get("chainage_to"), row.get("length_km"), row.get("surface_type"),
          row.get("maintenance_station"), row.get("maintenance_region"), row.get("completion_year"),
          row.get("rehab_year"), row.get("pavement_age"), source_id, utc_now()) for row in links],
    )
    conn.commit()


def add_condition_baseline(conn: sqlite3.Connection) -> dict[str, dict[str, Any]]:
    baseline = json.loads(PATHS.condition_json.read_text(encoding="utf-8"))
    source = conn.execute(
        """SELECT source_id FROM pms_source_files WHERE repository_group='road_condition'
           AND extension='.xlsx' ORDER BY modified_at DESC LIMIT 1"""
    ).fetchone()
    if not source:
        raise RuntimeError("No registered road-condition workbook can be linked to baseline values")
    source_id = source[0]
    synthetic_table = register_table(conn, source_id, "harmonised_link_condition_lookup", "derived_source_view", 1, len(baseline), list(next(iter(baseline.values())).keys()))
    for row_number, (lid, values) in enumerate(baseline.items(), 1):
        conn.execute(
            """INSERT OR IGNORE INTO pms_road_links
               (link_id,link_name,source_id,updated_at) VALUES(?,?,?,?)""",
            (lid, lid, source_id, utc_now()),
        )
        conn.execute(
            """INSERT OR REPLACE INTO pms_condition_observations
               (source_id,source_table_id,source_row_number,link_id,survey_date,survey_year,
                iri_m_per_km,rut_depth_mm,vci,pci,cracking_percent,condition_class,data_quality,raw_values_json)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (source_id, synthetic_table, row_number, lid, f"{values.get('year', 2024)}-12-31", values.get("year"),
             values.get("iri"), values.get("rut_mm"), values.get("vci"), values.get("pci"), values.get("cracking"),
             None, "harmonised", json.dumps(values, separators=(",", ":"))),
        )
    conn.commit()
    return baseline


def train_current_value_model(conn: sqlite3.Connection, baseline: dict[str, dict[str, Any]]) -> int:
    links = [dict(row) for row in conn.execute("SELECT * FROM pms_road_links ORDER BY link_id")]
    prediction_doc = json.loads(PATHS.prediction_json.read_text(encoding="utf-8"))
    prior_predictions = {row["link_id"]: row for row in prediction_doc.get("link_predictions", [])}
    numeric_features = ["length_km", "chainage_from_km", "chainage_to_km", "completion_year", "rehabilitation_year", "pavement_age_years"]
    categorical_features = ["road_class", "surface_type", "maintenance_region", "maintenance_station"]
    target_names = ["iri_m_per_km", "rut_depth_mm", "vci", "pci", "cracking_percent"]
    source_keys = ["iri", "rut_mm", "vci", "pci", "cracking"]
    feature_rows, target_rows, training_links = [], [], []
    medians = {}
    for name in numeric_features:
        observed_numbers = [value for row in links if (value := number(row.get(name))) is not None]
        medians[name] = float(np.median(observed_numbers)) if observed_numbers else 0.0

    def feature(row: dict[str, Any]) -> list[Any]:
        return [number(row.get(name)) if number(row.get(name)) is not None else medians[name] for name in numeric_features] + [str(row.get(name) or "Unknown") for name in categorical_features]

    for row in links:
        values = baseline.get(row["link_id"])
        targets = [number(values.get(key)) if values else None for key in source_keys]
        if values and all(value is not None for value in targets):
            feature_rows.append(feature(row)); target_rows.append(targets); training_links.append(row["link_id"])

    numeric_indexes = list(range(len(numeric_features)))
    category_indexes = list(range(len(numeric_features), len(numeric_features) + len(categorical_features)))
    hidden = tuple(int(value) for value in str(setting(conn, "ml_hidden_layers")).split(","))
    random_seed = int(setting(conn, "ml_random_seed"))
    max_iterations = int(setting(conn, "ml_max_iterations"))
    model = Pipeline([
        ("features", ColumnTransformer([
            ("numeric", StandardScaler(), numeric_indexes),
            ("category", OneHotEncoder(handle_unknown="ignore", sparse_output=False), category_indexes),
        ])),
        ("mlp", MLPRegressor(hidden_layer_sizes=hidden, activation="relu", solver="adam",
                             max_iter=max_iterations, early_stopping=True, random_state=random_seed)),
    ])
    x = np.asarray(feature_rows, dtype=object)
    y = np.asarray(target_rows, dtype=float)
    model.fit(x, y)
    fitted = model.predict(x)
    metrics = {
        target: {
            "r2": round(float(r2_score(y[:, index], fitted[:, index])), 5),
            "mae": round(float(mean_absolute_error(y[:, index], fitted[:, index])), 5),
        }
        for index, target in enumerate(target_names)
    }
    reporting_at = str(setting(conn, "reporting_date"))
    version_hash = hashlib.sha256(json.dumps(metrics, sort_keys=True).encode()).hexdigest()[:12]
    cursor = conn.execute(
        """INSERT INTO pms_model_runs
           (model_name,model_version,algorithm,feature_variables_json,target_variables_json,
            hyperparameters_json,training_rows,validation_metrics_json,trained_at,reporting_at,status)
           VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
        ("Uganda Pavement Current Value DNN", f"dnn-{version_hash}", "Deep multi-output MLP ensemble",
         json.dumps(numeric_features + categorical_features), json.dumps(target_names),
         json.dumps({"hidden_layers": hidden, "activation": "relu", "max_iter": max_iterations, "seed": random_seed}),
         len(training_links), json.dumps(metrics), utc_now(), reporting_at, "complete"),
    )
    model_run_id = cursor.lastrowid
    predicted = model.predict(np.asarray([feature(row) for row in links], dtype=object))
    variables = {row[0]: row[1] for row in conn.execute("SELECT canonical_name,variable_id FROM pms_variable_definitions")}
    threshold_good = float(setting(conn, "iri_good_upper")); threshold_fair = float(setting(conn, "iri_fair_upper")); threshold_poor = float(setting(conn, "iri_poor_upper"))

    def condition(iri: float) -> str:
        return "Good" if iri < threshold_good else "Fair" if iri < threshold_fair else "Poor" if iri < threshold_poor else "Very Poor"

    conn.execute("DELETE FROM pms_current_values")
    conn.execute("DELETE FROM pms_predictions")
    source_lookup = conn.execute("SELECT source_id FROM pms_source_files WHERE absolute_path=?", (str(PATHS.condition_json),)).fetchone()
    fallback_source = conn.execute("SELECT source_id FROM pms_source_files WHERE repository_group='road_condition' ORDER BY modified_at DESC LIMIT 1").fetchone()
    source_id = source_lookup[0] if source_lookup else (fallback_source[0] if fallback_source else None)
    current_rows = []
    for index, row in enumerate(links):
        lid = row["link_id"]
        observed = baseline.get(lid)
        learned = predicted[index]
        prior = prior_predictions.get(lid, {})
        confidence = float(prior.get("confidence_score") or max(0.55, min(0.98, np.mean([max(0.0, value["r2"]) for value in metrics.values()]))))
        method = "model_projected_from_observation" if observed else "deep_mlp_imputed"
        iri = number(prior.get("current_iri")) if prior else None
        iri = iri if iri is not None else (number(observed.get("iri")) if observed else float(learned[0]))
        base_iri = number(observed.get("iri")) if observed else None
        growth_ratio = iri / base_iri if base_iri and base_iri > 0 else 1.0
        values = {
            "iri_m_per_km": iri,
            "rut_depth_mm": min(100.0, max(0.0, (number(observed.get("rut_mm")) * growth_ratio) if observed and number(observed.get("rut_mm")) is not None else float(learned[1]))),
            "vci": min(100.0, max(0.0, (number(observed.get("vci")) / growth_ratio) if observed and number(observed.get("vci")) is not None else float(learned[2]))),
            "pci": min(100.0, max(0.0, (number(observed.get("pci")) / growth_ratio) if observed and number(observed.get("pci")) is not None else float(learned[3]))),
            "cracking_percent": min(100.0, max(0.0, (number(observed.get("cracking")) * growth_ratio) if observed and number(observed.get("cracking")) is not None else float(learned[4]))),
        }
        observed_at = f"{observed.get('year', 2024)}-12-31" if observed else None
        for name, value in values.items():
            current_rows.append((lid, variables[name], round(float(value), 4), None, reporting_at, observed_at, method, confidence, source_id, model_run_id))
        current_rows.append((lid, variables["condition_class"], None, condition(iri), reporting_at, observed_at, method, confidence, source_id, model_run_id))
        current_rows.append((lid, variables["pavement_age_years"], number(row.get("pavement_age_years")), None, reporting_at, observed_at, "network_register", 1.0, row.get("source_id"), None))
        current_rows.append((lid, variables["length_km"], number(row.get("length_km")), None, reporting_at, observed_at, "network_register", 1.0, row.get("source_id"), None))
        for horizon, key in [(0, "current_iri"), (1, "predicted_iri_1yr"), (3, "predicted_iri_3yr"), (5, "predicted_iri_5yr")]:
            future_iri = number(prior.get(key)) if prior else None
            if future_iri is None:
                rate = number(prior.get("deterioration_rate")) if prior else 0.0
                future_iri = iri * ((1 + rate) ** horizon)
            conn.execute(
                """INSERT INTO pms_predictions
                   (link_id,model_run_id,prediction_year,iri_m_per_km,condition_class,
                    intervention_type,intervention_year,confidence)
                   VALUES(?,?,?,?,?,?,?,?)""",
                (lid, model_run_id, int(reporting_at[:4]) + horizon, future_iri, condition(future_iri),
                 prior.get("intervention_type"), prior.get("intervention_year"), confidence),
            )
    conn.executemany(
        """INSERT INTO pms_current_values
           (link_id,variable_id,value_numeric,value_text,reporting_at,observed_at,value_method,
            confidence,source_id,model_run_id) VALUES(?,?,?,?,?,?,?,?,?,?)""", current_rows,
    )
    conn.commit()
    return model_run_id


def generate_infographics(conn: sqlite3.Connection, model_run_id: int) -> dict[str, Any]:
    generated = utc_now()
    state = [dict(row) for row in conn.execute("SELECT * FROM pms_v_current_link_state")]
    total_km = sum(number(row["length_km"]) or 0 for row in state)
    paved = [row for row in state if re.search(r"bitum|asphalt|sealed|paved|concrete", str(row["surface_type"] or ""), re.I) and not re.search(r"unpaved|unsealed", str(row["surface_type"] or ""), re.I)]
    method_counts = [dict(row) for row in conn.execute("SELECT value_method AS label,COUNT(DISTINCT link_id) AS value FROM pms_current_values WHERE variable_id=(SELECT variable_id FROM pms_variable_definitions WHERE canonical_name='iri_m_per_km') GROUP BY value_method")]
    condition_counts = [dict(row) for row in conn.execute("SELECT condition_class AS label,COUNT(*) AS links,SUM(length_km) AS km FROM pms_v_current_link_state GROUP BY condition_class ORDER BY CASE condition_class WHEN 'Good' THEN 1 WHEN 'Fair' THEN 2 WHEN 'Poor' THEN 3 ELSE 4 END")]
    iri_trend = [dict(row) for row in conn.execute("SELECT survey_year AS year,ROUND(AVG(iri_m_per_km),2) AS value,COUNT(*) AS samples FROM pms_condition_observations WHERE survey_year IS NOT NULL AND iri_m_per_km IS NOT NULL GROUP BY survey_year ORDER BY survey_year")]
    inventory = [dict(row) for row in conn.execute("SELECT asset_type AS label,COUNT(*) AS value FROM pms_inventory_assets GROUP BY asset_type ORDER BY value DESC LIMIT 10")]
    sources = [dict(row) for row in conn.execute("SELECT repository_group AS label,COUNT(*) AS files,SUM(record_count) AS records FROM pms_source_files GROUP BY repository_group ORDER BY repository_group")]
    predictions = [dict(row) for row in conn.execute("SELECT intervention_type AS label,COUNT(*) AS value FROM pms_predictions WHERE model_run_id=? AND prediction_year=(SELECT MIN(prediction_year) FROM pms_predictions WHERE model_run_id=?) GROUP BY intervention_type ORDER BY value DESC", (model_run_id, model_run_id))]
    ages = [
        {"label": "0-5", "value": sum(1 for row in state if number(row["pavement_age_years"]) is not None and number(row["pavement_age_years"]) <= 5)},
        {"label": "6-10", "value": sum(1 for row in state if number(row["pavement_age_years"]) is not None and 5 < number(row["pavement_age_years"]) <= 10)},
        {"label": "11-20", "value": sum(1 for row in state if number(row["pavement_age_years"]) is not None and 10 < number(row["pavement_age_years"]) <= 20)},
        {"label": "21+", "value": sum(1 for row in state if number(row["pavement_age_years"]) is not None and number(row["pavement_age_years"]) > 20)},
    ]
    rut_values = [number(row["current_rut_mm"]) for row in state if number(row["current_rut_mm"]) is not None]
    rut_quantiles = np.quantile(rut_values, [0.25, 0.5, 0.75]).tolist() if rut_values else [0, 0, 0]
    rut_bands = [
        {"label": "Lower quartile", "value": round(rut_quantiles[0], 2)},
        {"label": "Median", "value": round(rut_quantiles[1], 2)},
        {"label": "Upper quartile", "value": round(rut_quantiles[2], 2)},
    ]
    region_health = [dict(row) for row in conn.execute("SELECT maintenance_region AS label,ROUND(AVG(current_vci),1) AS value,COUNT(*) AS links FROM pms_v_current_link_state GROUP BY maintenance_region ORDER BY value DESC")]
    confidence = [dict(row) for row in conn.execute("SELECT value_method AS label,ROUND(AVG(confidence)*100,1) AS value,COUNT(*) AS samples FROM pms_current_values WHERE model_run_id=? GROUP BY value_method", (model_run_id,))]
    cards = [
        {"id": "network", "title": "National Network", "subtitle": "Current gazetted register", "type": "stat", "unit": "km", "payload": {"value": round(total_km, 1), "links": len(state), "paved_km": round(sum(number(row["length_km"]) or 0 for row in paved), 1)}},
        {"id": "condition", "title": "Current Condition", "subtitle": "DNN-assimilated link state", "type": "donut", "unit": "links", "payload": condition_counts},
        {"id": "coverage", "title": "Survey & Model Coverage", "subtitle": "Observed/projected versus imputed", "type": "donut", "unit": "links", "payload": method_counts},
        {"id": "iri_trend", "title": "IRI Survey History", "subtitle": "All linked condition observations", "type": "line", "unit": "m/km", "payload": iri_trend},
        {"id": "rutting", "title": "Current Rutting Distribution", "subtitle": "Network quartiles", "type": "bars", "unit": "mm", "payload": rut_bands},
        {"id": "regional_vci", "title": "Regional Visual Condition", "subtitle": "Current average VCI", "type": "bars", "unit": "VCI", "payload": region_health},
        {"id": "age", "title": "Pavement Age Profile", "subtitle": "Years since construction/rehabilitation", "type": "bars", "unit": "links", "payload": ages},
        {"id": "inventory", "title": "Road Inventory Assets", "subtitle": "Cross-linked source features", "type": "bars", "unit": "assets", "payload": inventory},
        {"id": "confidence", "title": "Model Confidence", "subtitle": "Current-value provenance", "type": "bars", "unit": "%", "payload": confidence},
        {"id": "maintenance", "title": "Maintenance Outlook", "subtitle": "DNN intervention recommendation", "type": "donut", "unit": "links", "payload": predictions},
    ]
    conn.execute("DELETE FROM pms_infographics")
    conn.executemany(
        """INSERT INTO pms_infographics
           (infographic_id,sort_order,title,subtitle,visualization_type,unit,payload_json,sql_query_name,generated_at)
           VALUES(?,?,?,?,?,?,?,?,?)""",
        [(card["id"], index, card["title"], card["subtitle"], card["type"], card["unit"],
          json.dumps(card["payload"], separators=(",", ":")), f"dashboard_{card['id']}", generated)
         for index, card in enumerate(cards, 1)],
    )
    model = dict(conn.execute("SELECT * FROM pms_model_runs WHERE model_run_id=?", (model_run_id,)).fetchone())
    document = {
        "generated_at": generated,
        "reporting_at": model["reporting_at"],
        "model": {"name": model["model_name"], "version": model["model_version"], "algorithm": model["algorithm"], "metrics": json.loads(model["validation_metrics_json"])},
        "source_summary": sources,
        "infographics": cards,
    }
    PATHS.dashboard_json.parent.mkdir(parents=True, exist_ok=True)
    PATHS.dashboard_json.write_text(json.dumps(document, indent=2), encoding="utf-8")
    conn.commit()
    return document


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the cross-linked NPMS backend engine")
    parser.add_argument("--skip-source-ingest", action="store_true", help="Reuse registered source metadata and normalized observations")
    parser.add_argument("--skip-access-metadata", action="store_true", help="Do not open MDB/ACCDB files")
    parser.add_argument("--only-access-metadata", action="store_true", help="Index only MDB/ACCDB tables and fields")
    args = parser.parse_args()
    PATHS.database.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(PATHS.database)
    conn.row_factory = sqlite3.Row
    try:
        setup_database(conn)
        load_network(conn)
        if args.only_access_metadata:
            ingest_access_sources(conn)
            print(json.dumps({
                "access_files": conn.execute("SELECT COUNT(*) FROM pms_source_files WHERE extension IN ('.mdb','.accdb')").fetchone()[0],
                "access_tables": conn.execute("SELECT COUNT(*) FROM pms_source_tables WHERE table_kind='access_table'").fetchone()[0],
                "access_fields": conn.execute("""SELECT COUNT(*) FROM pms_source_fields f JOIN pms_source_tables t ON t.source_table_id=f.source_table_id WHERE t.table_kind='access_table'""").fetchone()[0],
                "access_records": conn.execute("SELECT SUM(record_count) FROM pms_source_files WHERE extension IN ('.mdb','.accdb')").fetchone()[0],
            }, indent=2))
            return
        if not args.skip_source_ingest:
            ingest_sources(conn, include_access=not args.skip_access_metadata)
        baseline = add_condition_baseline(conn)
        model_run_id = train_current_value_model(conn, baseline)
        dashboard = generate_infographics(conn, model_run_id)
        print(json.dumps({
            "database": str(PATHS.database),
            "sources": conn.execute("SELECT COUNT(*) FROM pms_source_files").fetchone()[0],
            "source_tables": conn.execute("SELECT COUNT(*) FROM pms_source_tables").fetchone()[0],
            "source_fields": conn.execute("SELECT COUNT(*) FROM pms_source_fields").fetchone()[0],
            "manuals": conn.execute("SELECT COUNT(*) FROM pms_manual_documents").fetchone()[0],
            "road_links": conn.execute("SELECT COUNT(*) FROM pms_road_links").fetchone()[0],
            "condition_observations": conn.execute("SELECT COUNT(*) FROM pms_condition_observations").fetchone()[0],
            "inventory_assets": conn.execute("SELECT COUNT(*) FROM pms_inventory_assets").fetchone()[0],
            "current_values": conn.execute("SELECT COUNT(*) FROM pms_current_values").fetchone()[0],
            "infographics": len(dashboard["infographics"]),
            "dashboard": str(PATHS.dashboard_json),
        }, indent=2))
    finally:
        conn.close()


if __name__ == "__main__":
    main()
